#!/usr/bin/python
# coding=UTF-8


import run
import sys
from docx import Document
from docx.shared import Inches

document = Document()

# document.add_heading('Document Title', 0)
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')


def create(tableList):
    table = document.add_table(rows=0, cols=2, style='Table Grid')
    # table.style = 'Table Grid'
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = '資料表'
    hdr_cells[1].text = '說明'
    for key, value in tableList.items():
        row_cells = table.add_row().cells
        table_name_array = key.split(': ', 1)
        row_cells[0].text = str(table_name_array[0])
        row_cells[1].text = str(table_name_array[1])
    document.add_paragraph()

    for key, value in tableList.items():
        table = document.add_table(rows=0, cols=4, style='Table Grid')
        # table.style = 'Table Grid'
        row_cells = table.add_row().cells
        row = table.rows[0]
        a, b, c, d = row.cells[:4]
        a.merge(b).merge(c).merge(d)
        row_cells[0].text = str(key)
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = '資料行名稱'
        hdr_cells[1].text = '資料型別'
        hdr_cells[2].text = '主鍵'
        hdr_cells[3].text = '備註'

        for column in value:
            row_cells = table.add_row().cells
            row_cells[0].text = column["Field"]
            row_cells[1].text = column["Type"]
            row_cells[2].text = column["Key"]
            row_cells[3].text = column["Comment"]

        document.add_paragraph()

    # document.add_page_break()
    document.save('doc/' + run.dbName + '.docx')
