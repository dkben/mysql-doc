# 匯出 docx 的範例程式，可以直接跑這支測試效果

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam', 'aaa'),
    (7, '422', 'Eggs', 'bbb'),
    (4, '631', 'Spam, spam, eggs, and spam', 'ccc')
)

table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '資料行名稱'
hdr_cells[1].text = '資料型別說明'
hdr_cells[2].text = '主鍵'
hdr_cells[3].text = '備註'
row = table.rows[0]
a, b, c, d = row.cells[:4]
a.merge(b).merge(c).merge(d)
for qty, id, desc, desc2 in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = desc

document.add_page_break()

document.save('demo.docx')